"""
Report mimicry support (Feature 3)
─────────────────────────────────────────────────────────────────
Pipeline pieces for "write a report in the style of a sample":

  1. parse_sample()             sample report (DOCX / PDF / MD) -> outline
                                (section nodes with body text) + full text
  2. extract_material_context() per-material structure fingerprint (headings
                                + representative data lines + language)
  3. infer_rules_from_sample()  LLM derives ONE shared metric-rule set from
                                the sample; locator fields (section_hint /
                                aliases) merge the terminology of ALL
                                materials so the same rules match every file
                                even across languages
  4. save_as_profile()          optionally persist inferred rules as a new
                                profile (rules.xlsx + copied base config)
  5. write_report_docx()        per-section LLM writing that mimics the
                                sample's tone/structure using extracted
                                metric values (multi-entity aware)

The inferred rules are plain dicts compatible with engine.extract.extract_all
(rules=...), so they can be used directly without persisting a profile.
"""
from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

ROOT = Path(__file__).resolve().parent.parent
PROFILES_DIR = ROOT / "profiles"


# ════════════════════════════════════════════════════════════════════════
#  1. Sample parsing (DOCX / PDF / MD)
# ════════════════════════════════════════════════════════════════════════
@dataclass
class OutlineNode:
    """One section of the sample report."""
    level: int                  # 1=H1, 2=H2, ...
    title: str
    body: List[str] = field(default_factory=list)   # paragraphs (no sub-sections)

    def to_dict(self) -> Dict[str, Any]:
        return {"level": self.level, "title": self.title, "body": list(self.body)}


_HEADING_STYLE_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)


def _parse_docx(path: Path) -> List[OutlineNode]:
    """Read a DOCX via python-docx, splitting on Heading N styles. Falls back
    to a single section when no headings exist."""
    from docx import Document

    doc = Document(str(path))
    nodes: List[OutlineNode] = []
    current: Optional[OutlineNode] = None

    def _ensure_current():
        nonlocal current
        if current is None:
            current = OutlineNode(level=1, title="Body")
            nodes.append(current)

    def _iter_blocks(document):
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    for block in _iter_blocks(doc):
        cls = block.__class__.__name__
        if cls == "Paragraph":
            text = block.text.strip()
            if not text:
                continue
            style = (block.style.name or "") if block.style else ""
            m = _HEADING_STYLE_RE.match(style.strip())
            if m or style.lower().startswith("title"):
                level = int(m.group(1)) if m else 1
                current = OutlineNode(level=level, title=text)
                nodes.append(current)
            else:
                _ensure_current()
                current.body.append(text)
        elif cls == "Table":
            _ensure_current()
            rows = []
            for row in block.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                current.body.append("[table]\n" + "\n".join(rows))

    if not nodes:
        nodes.append(OutlineNode(level=1, title="Body"))
    return nodes


# physical-slice wire-format markers ("# --- PDF 物理切片 ... ---") are not
# document sections — they must never become outline nodes
_SLICE_MARKER_RE = re.compile(r"^\s*#\s*---.*---\s*$")


def _parse_markdown(text: str) -> List[OutlineNode]:
    """Build the outline from markdown headings (# .. ######)."""
    nodes: List[OutlineNode] = []
    current: Optional[OutlineNode] = None
    in_code = False

    for line in text.splitlines():
        if line.strip().startswith("```"):
            in_code = not in_code
            if current is not None:
                current.body.append(line)
            continue
        if in_code:
            if current is not None:
                current.body.append(line)
            continue
        if _SLICE_MARKER_RE.match(line):
            continue
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            current = OutlineNode(level=len(m.group(1)),
                                  title=m.group(2).strip())
            nodes.append(current)
        else:
            txt = line.rstrip()
            if not txt and (current is None or not current.body):
                continue
            if current is None:
                current = OutlineNode(level=1, title="Body")
                nodes.append(current)
            current.body.append(txt)

    if not nodes:
        nodes.append(OutlineNode(level=1, title="Body", body=[text]))
    return nodes


def parse_sample(
    path: str,
    *,
    profile=None,
    api_key: str = "",
    convert_engine: Optional[str] = None,
    log: Callable[[str], None] = print,
    should_stop: Callable[[], bool] = lambda: False,
) -> Tuple[List[OutlineNode], str]:
    """
    Parse the sample report. Returns (outline nodes, full plain text).
    Supports .docx / .pdf / .md / .txt; PDF samples go through the regular
    OCR pipeline first (engine.convert.pdf_to_md).
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Sample file not found: {path}")
    ext = p.suffix.lower()

    if ext == ".docx":
        log(f"[sample] parsing DOCX: {p.name}")
        nodes = _parse_docx(p)
    elif ext == ".pdf":
        if profile is None:
            raise ValueError("Parsing a PDF sample requires a profile "
                             "(for the docling/mineru engine choice)")
        from engine.convert import pdf_to_md
        log(f"[sample] PDF sample goes through OCR first "
            f"(engine={convert_engine or 'profile default'})")
        md_path = pdf_to_md(str(p), profile, api_key=api_key,
                            engine=convert_engine, log=log,
                            should_stop=should_stop)
        nodes = _parse_markdown(Path(md_path).read_text(encoding="utf-8"))
    elif ext in (".md", ".markdown", ".txt"):
        log(f"[sample] parsing text: {p.name}")
        nodes = _parse_markdown(p.read_text(encoding="utf-8"))
    else:
        raise ValueError(f"Unsupported sample format: {ext} "
                         f"(.docx / .pdf / .md / .txt)")

    lines: List[str] = []
    for n in nodes:
        lines.append(("#" * max(1, min(6, n.level))) + " " + n.title)
        lines.extend(n.body)
        lines.append("")
    return nodes, "\n".join(lines).strip()


# ════════════════════════════════════════════════════════════════════════
#  2. Language detection + material structure fingerprints
# ════════════════════════════════════════════════════════════════════════
def detect_language(text: str) -> str:
    """Rough zh / en detection by CJK-vs-Latin character counts."""
    if not text:
        return "zh"
    cjk = sum(1 for ch in text if "一" <= ch <= "鿿")
    latin = sum(1 for ch in text if "a" <= ch.lower() <= "z")
    if cjk == 0 and latin == 0:
        return "zh"
    return "zh" if cjk >= latin else "en"


_NUM_RE = re.compile(r"\d")


def extract_material_context(
    md_text: str,
    *,
    max_headings: int = 120,
    max_data_lines: int = 80,
    max_chars: int = 9000,
) -> Tuple[str, str]:
    """
    Build a "structure fingerprint" of one material MD: all headings plus
    representative number-bearing lines (table headers / financial rows).
    Returns (fingerprint text, detected language). Only the fingerprint is
    sent to the LLM — full documents would blow the context.
    """
    headings: List[str] = []
    data_lines: List[str] = []
    in_code = False
    for line in md_text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("```"):
            in_code = not in_code
            continue
        if s.startswith("#"):
            headings.append(s)
            continue
        if ("|" in s and _NUM_RE.search(s)) or \
           (_NUM_RE.search(s) and len(s) < 160 and sum(c.isdigit() for c in s) >= 3):
            data_lines.append(s)

    headings = headings[:max_headings]
    seen, uniq = set(), []
    for d in data_lines:
        key = d[:80]
        if key in seen:
            continue
        seen.add(key)
        uniq.append(d)
        if len(uniq) >= max_data_lines:
            break

    lang = detect_language(" ".join(headings) + " ".join(uniq[:30]))

    parts = ["[Material headings]"]
    parts.extend(headings or ["(no markdown headings found)"])
    parts.append("\n[Representative data lines]")
    parts.extend(uniq or ["(no obvious table/data lines found)"])
    fp = "\n".join(parts)
    if len(fp) > max_chars:
        fp = fp[:max_chars] + "\n...(fingerprint truncated)"
    return fp, lang


def extract_materials_context(
    materials: List[Tuple[str, str]],
    *,
    per_material_chars: int = 4500,
    max_chars: int = 14000,
) -> Tuple[str, str, List[str]]:
    """
    Multi-material fingerprint: label each (entity, md_text) block.
    Returns (combined fingerprint, primary language, per-material languages).
    """
    blocks: List[str] = []
    langs: List[str] = []
    for label, md in materials:
        fp, lang = extract_material_context(md, max_chars=per_material_chars)
        langs.append(lang)
        blocks.append(f"========== Material: {label} (language≈{lang}) ==========\n{fp}")
    combined = "\n\n".join(blocks)
    if len(combined) > max_chars:
        combined = combined[:max_chars] + "\n...(multi-material fingerprint truncated)"
    primary = max(set(langs), key=langs.count) if langs else "zh"
    return combined, primary, langs


# ════════════════════════════════════════════════════════════════════════
#  3. LLM rule inference
# ════════════════════════════════════════════════════════════════════════
_INFER_SYSTEM = """You are a methodology expert for financial/business analysis \
reports. Read the SAMPLE REPORT and list every QUANTITATIVE metric it cites \
(financial figures, KPIs, ratios, growth rates), so the metrics can later be \
extracted from raw source material by rule.

Output STRICT JSON only:
{
  "rules": [
    {
      "id": "M01",
      "name": "Revenue",            // metric name, in the sample's language
      "source": "Income statement - total operating revenue",  // one-line locator/definition
      "section_hint": ["Income Statement", "Financial Highlights"],  // section keywords likely found in source material
      "aliases": ["Revenue", "Total revenue"],   // synonym terms for text matching
      "group": "Profitability",
      "extraction_mode": "direct"   // direct, or "calc" for derived metrics
    }
  ]
}

Rules:
1. Quantitative metrics only — each must read as a number (incl. percentages). \
Skip narrative section titles.
2. List each metric once; merge repeats.
3. section_hint must be keywords plausible in SOURCE material, not the sample's headings.
4. Give several aliases per metric.
5. Derived metrics (e.g. ROE = net profit / avg equity) use extraction_mode \
"calc" with the formula explained in source.
6. Number ids M01 upward, in order of appearance.
7. Output the JSON object only — no prose, no fences."""


_INFER_SYSTEM_LOCALIZED = """You are a methodology expert for financial/business \
analysis reports. You receive two inputs:
(A) SAMPLE REPORT — decides WHICH quantitative metrics are needed.
(B) MATERIAL STRUCTURE FINGERPRINTS — the actual source documents the data \
will be extracted from. There may be SEVERAL materials (several entities), \
and their languages/terminology may differ from the sample and from each other \
(e.g. Chinese sample, English materials).

Your task: derive ONE unified metric list from (A) shared by all entities, but \
make every rule's locator fields (section_hint / aliases) MERGE the real \
terminology of ALL materials in (B), so the same rule set text-matches every \
material and the extracted rows align across entities.

Output STRICT JSON only:
{
  "material_language": "zh or en (dominant language of the materials)",
  "rules": [
    {
      "id": "M01",
      "name": "营业收入",            // metric name in the SAMPLE's language (used when writing the report)
      "name_in_material": "Revenue", // the metric's main name inside the materials
      "source": "Income statement - total operating revenue",  // one-line definition, sample language
      "section_hint": ["营业收入", "Income Statement", "合并利润表"],  // REAL heading keywords from the fingerprints, all languages merged
      "aliases": ["营业收入", "Revenue", "Total revenue"],            // REAL terms from the fingerprints, all languages merged
      "group": "Profitability",      // sample language
      "extraction_mode": "direct"
    }
  ]
}

Rules:
1. Quantitative metrics only. ONE unified list for all entities — never \
different metrics per material.
2. WHAT to extract comes from sample (A); section_hint and aliases MUST \
prioritize words that actually appear in the fingerprints (B) — copy heading \
wording verbatim where possible. Only add standard translations as fallback \
when a material lacks the term.
3. name and group use the sample's language; name_in_material, section_hint, \
aliases use the materials' languages (mixed is expected).
4. Derived metrics use extraction_mode "calc" with the formula in source.
5. Number ids M01 upward. Output the JSON object only — no prose, no fences."""


def infer_rules_from_sample(
    sample_text: str,
    *,
    api_key: str,
    base_url: str,
    text_model: str,
    materials: Optional[List[Tuple[str, str]]] = None,
    system_plain: Optional[str] = None,
    system_localized: Optional[str] = None,
    log: Callable[[str], None] = print,
    max_sample_chars: int = 30000,
) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    """
    LLM-infer the metric rules a sample report requires.
    Returns (rules list, dominant material language).

    materials: [(entity label, md text), ...] — when given, per-material
    structure fingerprints are sent along so the generated section_hint /
    aliases merge ALL materials' languages and terms (one shared rule set
    that matches every file). Without materials, plain sample-only inference.
    """
    from engine.llm import call_chat, safe_json_loads

    text = sample_text.strip()
    if len(text) > max_sample_chars:
        log(f"[infer] sample too long ({len(text)} chars), "
            f"truncated to {max_sample_chars}")
        text = text[:max_sample_chars]

    material_lang: Optional[str] = None
    materials = [m for m in (materials or []) if m and m[1] and m[1].strip()]
    if materials:
        fingerprint, material_lang, langs = extract_materials_context(materials)
        log(f"[infer] {len(materials)} material fingerprint(s) built "
            f"(languages: {langs}, primary={material_lang})")
        system = system_localized or _INFER_SYSTEM_LOCALIZED
        user = (
            f"[SAMPLE REPORT A — decides which metrics]\n```\n{text}\n```\n\n"
            f"[MATERIAL FINGERPRINTS B — {len(materials)} material(s), decide "
            f"locator terms; primary language≈{material_lang}]\n"
            f"```\n{fingerprint}\n```\n\n"
            f"Output the strict JSON (material_language + rules)."
        )
    else:
        system = system_plain or _INFER_SYSTEM
        user = (f"Sample report full text:\n\n```\n{text}\n```\n\n"
                f"Output the strict JSON.")

    raw = call_chat(
        [{"role": "system", "content": system},
         {"role": "user", "content": user}],
        model=text_model, api_key=api_key, base_url=base_url,
        temperature=0.0, timeout=300.0,
    )
    parsed = safe_json_loads(raw)
    if not isinstance(parsed, dict) or "rules" not in parsed:
        log(f"[infer] ⚠ LLM response missing 'rules': {raw[:400]}")
        raise ValueError("LLM did not return valid rules JSON — retry or "
                         "switch model")

    if parsed.get("material_language"):
        material_lang = str(parsed["material_language"]).strip().lower()

    norm: List[Dict[str, Any]] = []
    seen_ids = set()
    for i, r in enumerate(parsed["rules"] or [], 1):
        if not isinstance(r, dict):
            continue
        rid = str(r.get("id") or f"M{i:02d}").strip()
        if rid in seen_ids:
            rid = f"{rid}_{i}"
        seen_ids.add(rid)

        aliases = list(r.get("aliases") or [])
        nim = str(r.get("name_in_material") or "").strip()
        # the in-material name must participate in matching — fold into aliases
        if nim and nim not in aliases:
            aliases.insert(0, nim)

        norm.append({
            "id": rid,
            "name": str(r.get("name") or "").strip(),
            "name_in_material": nim,
            "source": str(r.get("source") or "").strip(),
            "section_hint": list(r.get("section_hint") or []),
            "aliases": aliases,
            "group": str(r.get("group") or "inferred").strip(),
            "extraction_mode": (r.get("extraction_mode") or "direct").strip(),
            "enabled": True,
        })
    log(f"[infer] {len(norm)} metric rule(s) inferred"
        + (f" (material language={material_lang})" if material_lang else ""))
    return norm, material_lang


# ════════════════════════════════════════════════════════════════════════
#  4. Persist inferred rules as a new profile (optional)
# ════════════════════════════════════════════════════════════════════════
_SAFE_NAME_RE = re.compile(r"[^a-zA-Z0-9_\-]")


def _sanitize_profile_name(name: str) -> str:
    name = name.strip().replace(" ", "_")
    name = _SAFE_NAME_RE.sub("_", name)
    return name or "inferred"


def _write_rules_xlsx(rules: List[Dict[str, Any]], xlsx_path: Path):
    """Write the rules into an Excel the engine's rules_excel can read back."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "rules"
    headers = ["id", "name", "name_in_material", "source", "section_hint",
               "aliases", "extraction_mode", "enabled", "group"]
    ws.append(headers)
    for r in rules:
        ws.append([
            r.get("id"),
            r.get("name"),
            r.get("name_in_material", ""),
            r.get("source", ""),
            ", ".join(r.get("section_hint") or []),
            ", ".join(r.get("aliases") or []),
            r.get("extraction_mode", "direct"),
            "TRUE" if r.get("enabled", True) else "FALSE",
            r.get("group", ""),
        ])
    for i, w in enumerate([8, 22, 22, 50, 40, 40, 16, 10, 18], start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(xlsx_path))


def save_as_profile(
    rules: List[Dict[str, Any]],
    profile_name: str,
    *,
    base_profile: str = "cn_securities",
    display_name: str = "",
    language: str = "zh",
    overwrite: bool = False,
    log: Callable[[str], None] = print,
) -> str:
    """
    Persist the inferred rules as profiles/<profile_name>/ (rules.xlsx plus
    the base profile's profile.json + prompts/). Returns the saved name.
    """
    safe = _sanitize_profile_name(profile_name)
    out_dir = PROFILES_DIR / safe
    if out_dir.exists() and not overwrite:
        raise FileExistsError(f"Profile already exists: {out_dir} "
                              f"(pass overwrite to replace)")

    base_dir = PROFILES_DIR / base_profile
    if not base_dir.exists():
        raise FileNotFoundError(f"Base profile not found: {base_profile}")

    out_dir.mkdir(parents=True, exist_ok=True)
    base_cfg = json.loads((base_dir / "profile.json").read_text(encoding="utf-8"))
    base_cfg["display_name"] = display_name or f"Inferred - {profile_name}"
    base_cfg["language"] = language
    base_cfg["rules_file"] = "rules/rules.xlsx"
    (out_dir / "profile.json").write_text(
        json.dumps(base_cfg, ensure_ascii=False, indent=2), encoding="utf-8")

    src_prompts = base_dir / "prompts"
    if src_prompts.exists():
        dst_prompts = out_dir / "prompts"
        dst_prompts.mkdir(parents=True, exist_ok=True)
        for f in src_prompts.iterdir():
            if f.is_file():
                shutil.copyfile(f, dst_prompts / f.name)

    _write_rules_xlsx(rules, out_dir / "rules" / "rules.xlsx")
    log(f"[profile] saved: {out_dir} ({len(rules)} rules, base={base_profile})")
    return safe


# ════════════════════════════════════════════════════════════════════════
#  5. Style mimicry + Word writing
# ════════════════════════════════════════════════════════════════════════
_WRITE_SYSTEM = """You are a professional report writer. You receive (a) one \
section's SAMPLE TEXT (to learn tone/structure/terminology), (b) the DATA \
METRICS relevant to that section, and (c) the section title. Rewrite the \
section body in the sample's style, using the provided data.

Strict requirements:
1. Output plain body text only — do not repeat the title; separate paragraphs \
with blank lines.
2. Use the provided data; NEVER invent numbers. For null/missing values write \
"not disclosed" (or the natural equivalent in the writing language) or skip \
gracefully.
3. Keep the sample's sentence rhythm, modifier density, and terminology — but \
do not copy the sample verbatim (its numbers belong to another entity/period).
4. You may compute and narrate changes (YoY, deltas) from the given values.
5. If the data covers MULTIPLE entities (benchmarking), write a comparative \
narrative: rankings, gaps, leaders/laggards — not an entity-by-entity list.
6. Write in the same language as the sample text.
7. No markdown heading symbols (#); no bullet markers at paragraph starts.
8. 1-4 paragraphs, similar to the sample section's length."""

DEFAULT_WRITE_PROMPT = _WRITE_SYSTEM
DEFAULT_INFER_PROMPT = _INFER_SYSTEM
DEFAULT_INFER_PROMPT_LOCALIZED = _INFER_SYSTEM_LOCALIZED


def _node_match_metrics(
    node: OutlineNode, metric_records: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """Pick the metrics relevant to a section — simple containment matching
    of metric name/aliases/group against the section's title+body."""
    text = (node.title + "\n" + "\n".join(node.body)).lower()
    out = []
    for rec in metric_records:
        keys = [rec.get("name", ""), rec.get("name_in_material", ""),
                rec.get("group", "")]
        keys.extend(rec.get("aliases") or [])
        if any(k and str(k).lower() in text for k in keys):
            out.append(rec)
    return out


def _build_metric_block(records: List[Dict[str, Any]]) -> str:
    """Render metrics for the prompt. Two record shapes:
    single entity: {name, value, unit}
    multi entity:  {name, values: {entity: {value, unit}}}"""
    if not records:
        return ("(no directly relevant metrics — write a structural summary "
                "in the sample's style)")
    lines = []
    for r in records:
        name = r.get("name")
        if isinstance(r.get("values"), dict):          # multi-entity
            parts = []
            for ent, v in r["values"].items():
                val, unit = v.get("value"), v.get("unit") or ""
                parts.append(f"{ent}=missing" if val in (None, "", "null")
                             else f"{ent}={val}{unit}")
            lines.append(f"- {name}: " + "; ".join(parts))
        else:                                           # single entity
            val, unit = r.get("value"), r.get("unit") or ""
            lines.append(f"- {name}: missing" if val in (None, "", "null")
                         else f"- {name}: {val}{(' ' + unit) if unit else ''}")
    return "\n".join(lines)


def _write_section(
    node: OutlineNode,
    related: List[Dict[str, Any]],
    all_metrics: List[Dict[str, Any]],
    *,
    api_key: str, base_url: str, text_model: str,
    temperature: float = 0.3,
    system_prompt: Optional[str] = None,
    log: Callable[[str], None] = print,
) -> str:
    from engine.llm import call_chat

    sample_excerpt = "\n".join(node.body)[:4000]
    metric_block = _build_metric_block(related or all_metrics[:8])
    user = (
        f"[Section title] {node.title}\n\n"
        f"[Sample section text — learn the style]\n"
        f"{sample_excerpt or '(short sample section — use industry convention)'}\n\n"
        f"[Available data]\n{metric_block}\n\n"
        f"Rewrite this section body in the sample's style."
    )
    try:
        raw = call_chat(
            [{"role": "system", "content": system_prompt or _WRITE_SYSTEM},
             {"role": "user", "content": user}],
            model=text_model, api_key=api_key, base_url=base_url,
            temperature=temperature, timeout=240.0,
        )
        out = re.sub(r"^```.*?$", "", raw, flags=re.M).strip()
        out = re.sub(r"^\s*#{1,6}\s+.*$", "", out, flags=re.M).strip()
        return out
    except Exception as e:
        log(f"[write] ⚠ section '{node.title}' failed: {e}")
        return f"(section generation failed: {e})"


def write_report_docx(
    outline: List[OutlineNode],
    metric_records: List[Dict[str, Any]],
    *,
    out_path: str,
    api_key: str,
    base_url: str,
    text_model: str,
    title: str = "",
    write_system: Optional[str] = None,
    log: Callable[[str], None] = print,
    progress: Callable[[float], None] = lambda p: None,
    should_stop: Callable[[], bool] = lambda: False,
) -> str:
    """Generate the .docx along the sample outline; each section's body is
    LLM-written, assembled via python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    total = max(1, len(outline))
    for idx, node in enumerate(outline, 1):
        if should_stop():
            log("[write] abort signal received, stopping")
            break
        log(f"[write] [{idx}/{total}] section: {node.title}")
        related = _node_match_metrics(node, metric_records)
        doc.add_heading(node.title, level=max(1, min(9, node.level)))

        body = _write_section(
            node, related, metric_records,
            api_key=api_key, base_url=base_url, text_model=text_model,
            system_prompt=write_system, log=log,
        )
        for para in body.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(11)
        progress(idx / total)

    out_p = Path(out_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    log(f"[write] ✅ report written: {out_p}")
    return str(out_p)


# ════════════════════════════════════════════════════════════════════════
#  Multi-entity helpers (benchmarking)
# ════════════════════════════════════════════════════════════════════════
def metrics_records_from_df(df) -> List[Dict[str, Any]]:
    """extract_all DataFrame -> single-entity metric records."""
    out = []
    for _, row in df.iterrows():
        out.append({
            "id": row.get("id"),
            "name": row.get("name"),
            "group": row.get("group", ""),
            "value": row.get("value"),
            "unit": row.get("unit"),
            "aliases": [],
        })
    return out


def build_metric_records_multi(
    entity_results: Dict[str, Any],
    rule_aliases: Optional[Dict[str, List[str]]] = None,
) -> List[Dict[str, Any]]:
    """Merge per-entity DataFrames into metric records carrying
    values:{entity:{value,unit}}; order follows the first entity."""
    rule_aliases = rule_aliases or {}
    records: Dict[str, Dict[str, Any]] = {}
    order: List[str] = []
    for entity, df in entity_results.items():
        for _, row in df.iterrows():
            rid = row.get("id")
            if rid not in records:
                records[rid] = {
                    "id": rid,
                    "name": row.get("name"),
                    "group": row.get("group", ""),
                    "aliases": rule_aliases.get(rid, []),
                    "values": {},
                }
                order.append(rid)
            records[rid]["values"][entity] = {
                "value": row.get("value"),
                "unit": row.get("unit"),
            }
    return [records[r] for r in order]
