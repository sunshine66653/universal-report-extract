"""
Recognized MD -> Word (.docx)
─────────────────────────────────────────────────────────────────
Converts the engine's Markdown output (HTML <table> blocks + pipe tables) to
a Word document:
  - # / ## / ### ...  -> Word Heading 1-6
  - HTML <table> (rowspan/colspan expanded into a plain grid) -> Word table
  - | pipe tables |   -> Word table
  - physical-slice marker lines -> skipped (not part of the document)
  - ![](data:..) / ![](images/..) image refs -> dropped
Dependency: python-docx
"""
from __future__ import annotations

import argparse
import html as _html
import re
import sys
from pathlib import Path

# Windows GBK console compatibility: switch stdout/stderr to UTF-8
for _s in (sys.stdout, sys.stderr):
    if hasattr(_s, "reconfigure"):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

# Chinese wire-format marker written by engine.mineru — keep the pattern verbatim
_SLICE_RE = re.compile(
    r"^\s*#\s*---\s*PDF\s*物理切片\s*[:：]\s*第\s*\d+\s*-\s*\d+\s*页\s*---\s*$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_IMG_RE = re.compile(r"!\[[^\]]*\]\([^)]+\)")
_PIPE_ROW_RE = re.compile(r"^\s*\|.+\|\s*$")
_PIPE_SEP_RE = re.compile(r"^\s*\|?[\s:\-|]+\|?\s*$")


def _strip_inline_md(text: str) -> str:
    text = _IMG_RE.sub("", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    return text.strip()


# ── HTML table parsing (rowspan/colspan expanded into a plain grid) ─────
def _parse_html_table(table_html: str) -> list[list[str]]:
    rows_html = re.findall(r"<tr[^>]*>(.*?)</tr>", table_html,
                           re.DOTALL | re.IGNORECASE)
    grid: list[list[str | None]] = []
    spans: dict[tuple[int, int], str] = {}   # cells occupied by a rowspan above

    for ri, row_html in enumerate(rows_html):
        cells = re.findall(
            r"<(th|td)([^>]*)>(.*?)</\1>", row_html, re.DOTALL | re.IGNORECASE)
        grid.append([])
        ci = 0
        for _tag, attrs, content in cells:
            # skip grid positions claimed by a rowspan from a previous row
            while (ri, ci) in spans:
                grid[ri].append(spans.pop((ri, ci)))
                ci += 1
            text = re.sub(r"<[^>]+>", " ", content)
            text = _html.unescape(text)
            text = re.sub(r"\s+", " ", text).strip()

            cs = re.search(r'colspan\s*=\s*["\']?(\d+)', attrs, re.IGNORECASE)
            rs = re.search(r'rowspan\s*=\s*["\']?(\d+)', attrs, re.IGNORECASE)
            n_cs = int(cs.group(1)) if cs else 1
            n_rs = int(rs.group(1)) if rs else 1

            for k in range(n_cs):
                grid[ri].append(text if k == 0 else "")
                for r2 in range(1, n_rs):
                    spans[(ri + r2, ci)] = text if k == 0 else ""
                ci += 1
        # trailing rowspan leftovers at the end of the row
        while (ri, ci) in spans:
            grid[ri].append(spans.pop((ri, ci)))
            ci += 1

    width = max((len(r) for r in grid), default=0)
    return [[(c or "") for c in row] + [""] * (width - len(row))
            for row in grid if row]


def _parse_pipe_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for ln in lines:
        if _PIPE_SEP_RE.match(ln):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append([_strip_inline_md(c) for c in cells])
    width = max((len(r) for r in rows), default=0)
    return [r + [""] * (width - len(r)) for r in rows]


def _add_table(doc, rows: list[list[str]]):
    from docx.shared import Pt
    if not rows or not rows[0]:
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.font.bold = True
    doc.add_paragraph()


def md_to_docx(md_path: str | Path, out_path: str | Path,
               title: str = "") -> Path:
    from docx import Document
    from docx.shared import Pt

    md_path = Path(md_path)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    text = md_path.read_text(encoding="utf-8-sig", errors="replace")
    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    lines = text.splitlines()
    i = 0
    para_buf: list[str] = []

    def _flush_para():
        nonlocal para_buf
        if para_buf:
            content = _strip_inline_md(" ".join(para_buf))
            if content:
                p = doc.add_paragraph(content)
                for run in p.runs:
                    run.font.size = Pt(10.5)
            para_buf = []

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if _SLICE_RE.match(s):                       # slice marker: skip
            _flush_para(); i += 1; continue

        if not s:
            _flush_para(); i += 1; continue

        m = _HEADING_RE.match(s)
        if m:
            _flush_para()
            level = min(len(m.group(1)), 6)
            doc.add_heading(_strip_inline_md(m.group(2)), level=level)
            i += 1; continue

        low = s.lower()
        if low.startswith("<table"):                  # HTML table block
            _flush_para()
            block = [line]
            while "</table>" not in lines[i].lower() and i + 1 < len(lines):
                i += 1
                block.append(lines[i])
            _add_table(doc, _parse_html_table("\n".join(block)))
            i += 1; continue

        if _PIPE_ROW_RE.match(s):                     # pipe table block
            _flush_para()
            block = []
            while i < len(lines) and _PIPE_ROW_RE.match(lines[i].strip()):
                block.append(lines[i]); i += 1
            _add_table(doc, _parse_pipe_table(block))
            continue

        para_buf.append(s)
        i += 1

    _flush_para()
    doc.save(str(out_path))
    print(f"[docx] ✅ Word document written: {out_path}")
    return out_path


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Recognized MD -> Word (.docx)")
    ap.add_argument("md", help="recognized MD path")
    ap.add_argument("output", help="output .docx path")
    ap.add_argument("--title", default="", help="document title")
    args = ap.parse_args()
    md_to_docx(args.md, args.output, title=args.title)
